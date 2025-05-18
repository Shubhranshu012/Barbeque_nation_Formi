import os
import re
import json
from docx import Document
import google.generativeai as genai

# Initialize the Gemini client
genai.configure(api_key="AIzaSyD6fIAuoFd5TvbT1gBAse2Lx--Lg54Xvvc")

# Folder containing .docx files
docs_folder = 'docs'
parsed_folder = 'Parsed'
os.makedirs(parsed_folder, exist_ok=True)

def extract_json_from_response(text):
    # Remove code block formatting
    text = re.sub(r"```json|```", "", text).strip()

    # Extract first JSON object using regex
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        json_str = match.group(0)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            return None
    else:
        print("No JSON object found in response.")
        return None

# 💾 Save JSON to file based on outlet name
def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|]', '_', name)

# 💾 Save JSON to file in Parsed folder
def save_json_to_file(json_data, original_filename):
    # Use the original filename but with .json extension
    base_name = os.path.splitext(original_filename)[0]  # Removes .docx
    file_name = f"{sanitize_filename(base_name)}.json"
    file_path = os.path.join(parsed_folder, file_name)

    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, indent=4)
    print(f"✅ Saved JSON to: {file_path}")

# Extracts text from paragraphs and tables in a DOCX file
def parse_docx(file_path):
    doc = Document(file_path)
    full_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract tables (tab-separated)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append('\t'.join(row_text))

    return '\n'.join(full_text)

# Builds a structured prompt for Gemini
def build_prompt(raw_text):
    return f"""
You are a smart assistant. Given the following restaurant info, extract structured JSON with:
- city
- outlet
- address
- Weekly Instruction (flexible or Take till Last Enty or Till Some time)
- timings per day (Lunch & Dinner): open, last_entry, close

Here is the raw text:
\"\"\"{raw_text}\"\"\"
"""

# Main execution logic
def main():
    model = genai.GenerativeModel('gemini-1.5-flash')  # Or 'gemini-1.5-pro' if available

    for filename in os.listdir(docs_folder):
        if filename.endswith('.docx'):
            file_path = os.path.join(docs_folder, filename)
            raw_text = parse_docx(file_path)
            prompt = build_prompt(raw_text)

            response = model.generate_content(prompt)
            json_data = extract_json_from_response(response.text)
            if json_data:
                outlet_name = json_data.get('outlet', 'Unknown_Outlet')
                save_json_to_file(json_data, filename)
            else:
                print("❌ Failed to extract JSON.")

    print("\nAll files processed.")

if __name__ == "__main__":
    main()
